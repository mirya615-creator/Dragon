from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


OUTPUT = Path(r"D:\Codex project\Dragon\Docs\DragonBound\Analytics\DragonBound_Firebase_Analytics_Closed_Loop_Task_Plan_V1_2.docx")

NAVY = "17365D"
BLUE = "2E74B5"
PALE_BLUE = "E8EEF5"
LIGHT = "F2F4F7"
INK = "1F2937"
MUTED = "667085"
WHITE = "FFFFFF"
GREEN = "EAF4EA"
GOLD = "FFF4CE"
RED = "FDECEC"
FONT = "Microsoft YaHei"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths_dxa):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            set_cell_margins(cell)


def set_repeat_keep(paragraph, keep_next=False, keep_lines=True):
    p_pr = paragraph._p.get_or_add_pPr()
    if keep_next:
        p_pr.append(OxmlElement("w:keepNext"))
    if keep_lines:
        p_pr.append(OxmlElement("w:keepLines"))


def style_run(run, size=10.5, bold=False, color=INK, italic=False):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def add_para(doc, text="", *, bold_prefix=None, style=None, size=10.5, color=INK,
             align=None, before=0, after=6, italic=False):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.22
    if align is not None:
        p.alignment = align
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        style_run(r1, size=size, bold=True, color=color)
        r2 = p.add_run(text[len(bold_prefix):])
        style_run(r2, size=size, color=color, italic=italic)
    else:
        r = p.add_run(text)
        style_run(r, size=size, color=color, italic=italic)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.left_indent = Inches(0.375 if level == 0 else 0.625)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.2
    style_run(p.add_run(text), size=10.2)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.2
    style_run(p.add_run(text), size=10.2)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt({1: 18, 2: 14, 3: 10}[level])
    p.paragraph_format.space_after = Pt({1: 10, 2: 7, 3: 5}[level])
    style_run(p.add_run(text), size={1: 16, 2: 13, 3: 11.5}[level], bold=True,
              color=BLUE if level < 3 else NAVY)
    set_repeat_keep(p, keep_next=True)
    return p


def add_callout(doc, label, text, fill=PALE_BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    style_run(p.add_run(label + "  "), size=10.5, bold=True, color=NAVY)
    style_run(p.add_run(text), size=10.3, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc, headers, rows, widths, font_size=8.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_repeat_table_header(table.rows[0])
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, PALE_BLUE)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        style_run(p.add_run(header), size=font_size, bold=True, color=NAVY)
    for ridx, values in enumerate(rows):
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            cell = cells[idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if ridx % 2 == 1:
                set_cell_shading(cell, "FAFBFC")
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.08
            style_run(p.add_run(str(value)), size=font_size, color=INK)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_page_break(doc):
    doc.add_page_break()


def set_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    style_run(paragraph.add_run("DragonBound 埋点闭环任务书  |  "), size=8.5, color=MUTED)
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])
    style_run(run, size=8.5, color=MUTED)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.22

    for level, size in ((1, 16), (2, 13), (3, 11.5)):
        st = doc.styles[f"Heading {level}"]
        st.font.name = FONT
        st._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(BLUE if level < 3 else NAVY)

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    style_run(hp.add_run("DRAKEFORGE / ANALYTICS DELIVERY PLAN"), size=8.2, bold=True, color=MUTED)
    set_page_number(section.footer.paragraphs[0])


def build_doc():
    doc = Document()
    configure_document(doc)

    # Cover - editorial_cover pattern.
    add_para(doc, "DRAKEFORGE · PRODUCT / ENGINEERING / DATA", size=10, color=BLUE,
             align=WD_ALIGN_PARAGRAPH.CENTER, before=50, after=22)
    add_para(doc, "DragonBound Firebase\n完整埋点闭环任务书", size=27, color=NAVY,
             align=WD_ALIGN_PARAGRAPH.CENTER, after=12)
    add_para(doc, "客户端 × 服务端 × Firebase 三方实施与验收方案", size=14, color=BLUE,
             align=WD_ALIGN_PARAGRAPH.CENTER, after=36)
    add_callout(doc, "目标", "建立从事件定义、客户端采集、服务端权威确认、Firebase 分析、数据质量监控到产品决策的完整闭环。", fill=PALE_BLUE)
    add_para(doc, "版本：V1.2（增加客户端/服务端连接前后实施排期）", size=10.5, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, before=34, after=4)
    add_para(doc, "日期：2026-08-27", size=10.5, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    add_para(doc, "适用范围：Android / Unity 2022.3 / Firebase Analytics / GA4 / BigQuery（可选）",
             size=10, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    add_para(doc, "状态：供产品、客户端、服务端、数据与 QA 联合评审",
             size=10, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    add_page_break(doc)

    add_heading(doc, "0. 文档结论与执行摘要", 1)
    add_callout(doc, "结论", "完整闭环由客户端、服务端和 Firebase 三方即可完成。客户端负责玩法与交互观察；服务端负责资产、结算、排行和交易的权威结果；Firebase 负责接收、调试、分析和导出。前提是三方使用同一事件协议、单一上报归属和可追踪的关联键。", fill=GREEN)
    add_para(doc, "本任务书建议将终局事件统一为 run_end、波次结束统一为 wave_end。当前代码中的 match_finish、wave_finish 作为 V2 兼容名保留一个版本，进入 V3 后停止双写，避免同一事实被统计两次。", size=10.5)
    add_table(doc,
              ["阶段", "核心产出", "牵头方", "完成标志"],
              [
                  ("P0 协议冻结", "事件字典、字段、原因码、KPI、Build Lane", "产品/数据", "三方签字冻结 V3"),
                  ("P1 客户端基座", "Firebase 初始化、统一 Session、Firebase Sink", "客户端", "DebugView 收到连续事件"),
                  ("P2 玩法接线", "对局、波次、招募、英雄、Boss、道具、符文", "客户端", "核心漏斗事件完整"),
                  ("P3 权威业务", "体力、金币、奖励、商店、排行、账本", "服务端", "权威结果可对账"),
                  ("P4 数据运营", "GA4 探索、KPI 面板、BigQuery、告警", "Firebase/数据", "日报与质量告警运行"),
                  ("P5 发布验收", "QA、灰度、生产监控、回滚", "三方", "生产 Lane 稳定 7 天"),
              ], [1250, 3600, 1450, 3060], font_size=8.7)

    add_heading(doc, "1. 总体任务", 1)
    add_heading(doc, "1.1 建设目标", 2)
    for text in [
        "建立一套可版本化、可验证、可扩展的统一事件协议。",
        "实现客户端玩法事件实时采集，并保证失败不影响游戏主流程。",
        "让服务端权威结果可与客户端请求、对局和 Firebase 事件关联。",
        "在 Firebase DebugView、GA4 探索和 BigQuery 中完成从采集到分析的闭环。",
        "形成关键 KPI、数据质量指标、异常告警、变更评审和发布流程。",
    ]:
        add_bullet(doc, text)

    add_heading(doc, "1.2 闭环架构与单一事实来源", 2)
    add_para(doc, "客户端玩法运行时 → 类型化业务回调 → AnalyticsRunSession → AnalyticsRecorderV3 → FirebaseAnalyticsSink → Firebase/GA4 → BigQuery/报表。服务端通过业务响应携带 operation_id、server_event_id、ledger_status 等权威字段；客户端在收到响应后记录 server_confirmed 事件。")
    add_table(doc,
              ["数据类别", "事实来源", "主上报方", "说明"],
              [
                  ("玩法行为", "客户端运行时", "客户端", "波次、战斗、招募、阵型、英雄、Boss"),
                  ("UI/请求行为", "客户端 UI/Service Wrapper", "客户端", "点击、请求、SDK 返回；不代表奖励到账"),
                  ("资产与结算", "服务端账本", "客户端记录服务端确认；服务端留审计日志", "体力、金币、道具、符文、商店、广告奖励"),
                  ("排行", "服务端排行服务", "客户端记录服务端快照/变更", "服务端值为最终权威"),
                  ("分析与展示", "Firebase/GA4/BigQuery", "Firebase", "不作为资产账本"),
              ], [1600, 1900, 2400, 3460], font_size=8.5)
    add_callout(doc, "单写原则", "同一业务事实只能有一个 Analytics 主事件。客户端请求和服务端结果必须使用不同事件或不同 phase 字段，不能双方各发一条同名成功事件。", fill=GOLD)

    add_heading(doc, "1.3 总体完成定义（Definition of Done）", 2)
    for text in [
        "核心事件 Schema、触发点、字段、原因码、Owner 和 KPI 均有版本化文档。",
        "客户端真机 DebugView 能看到 run_start → wave_start/wave_end → run_end 的完整有序链路。",
        "所有服务端权威事件都有 operation_id、server_event_id 或哈希幂等引用，可与客户端请求关联。",
        "dev、qa、staging、production 数据严格隔离；生产报表默认只包含 production。",
        "事件丢失率、无效事件率、重复率、终局缺失率达到验收阈值。",
        "关键 KPI 在 GA4 或 BigQuery 中可复算，并与服务端账本抽样对账。",
        "用户同意、隐私政策、数据保留和访问权限完成审核。",
    ]:
        add_bullet(doc, text)

    add_page_break(doc)
    add_heading(doc, "2. 客户端具体任务步骤", 1)
    client_steps = [
        ("C1 项目与 Firebase 配置", "将 Android Package Name 固定为 com.drakeforge.mergedefense；Minimum API Level ≥ 23；导入同版本 FirebaseAnalytics 包；google-services.json 放入 Assets 根目录；执行 Android Resolver。", "Android 构建成功，依赖无冲突"),
        ("C2 Firebase 初始化", "在首场景调用 CheckAndFixDependenciesAsync；建立 Ready/Failed 状态；初始化完成前有界缓存；后台与退出时刷新。", "真机初始化 Available"),
        ("C3 Firebase Sink", "实现 IAnalyticsSinkV3/FirebaseAnalyticsSink；仅发送适用字段；异常隔离；提供写入失败计数。", "单元测试覆盖参数映射"),
        ("C4 统一 Run Session", "统一生成 run_id、event_id、sequence、build_lane、execution_context、config/build version；所有 Adapter 共享序列分配器。", "消除多 Adapter OutOfOrder"),
        ("C5 生命周期回调", "在 StartRun、BeginWave、EndWave、终局结算处公开类型化事件；禁止解析 Debug.Log。", "每局起止各一次、每波起止配对"),
        ("C6 玩法适配器", "接入敌人、Boss、招募、阵型、英雄形成/经验/升级、最后一击、道具、符文事件。", "核心事件全部有 EditMode 测试"),
        ("C7 服务端响应适配", "封装体力、金币、商店、排行和账本 API；仅在收到服务端确认后记录权威结果。广告暂不接入生产埋点。", "请求与结果可由 operation_id 关联"),
        ("C7A 广告观察口预留", "冻结 ad_request/ad_impression/ad_result/ad_reward_result 协议；若广告业务已存在，则增加 IAdAnalyticsObserver + NoOp；当前不接 Firebase、不发送数据。", "默认 NoOp、业务无副作用"),
        ("C8 同意与隐私", "提供 Analytics Collection 开关；未同意不采集；禁止发送名称、账号、设备、Token、聊天、原始交易号。", "隐私测试与政策通过"),
        ("C9 可靠性", "限制内存队列；记录丢弃数；不无限重试；Analytics 失败不得改变 RNG、结算和存档。", "故障注入测试通过"),
        ("C10 发布验证", "dev/qa/staging/prod Lane 打包验证；DebugView、Release IL2CPP/ARM64、断网/切后台/崩溃恢复测试。", "发布清单签字"),
    ]
    add_table(doc, ["步骤", "具体任务", "验收产物"], client_steps, [1500, 5260, 2600], font_size=8.3)

    add_heading(doc, "2.1 客户端代码结构建议", 2)
    for text in [
        "AnalyticsRunSession：唯一会话上下文和全局 sequence/event_id 分配。",
        "AnalyticsRecorderV3：Schema 校验、去重、顺序、Lane/Context 一致性检查。",
        "FirebaseAnalyticsSink：Firebase 参数映射、初始化状态、错误隔离。",
        "GameplayAnalyticsAdapter：对局、波次、敌人、Boss、英雄和阵型。",
        "EconomyAnalyticsAdapter：服务端确认的体力、金币、奖励、商店和账本。",
        "AnalyticsConsentProvider：采集许可与运行期开关。",
        "AnalyticsDiagnostics：只输出计数和稳定错误码，不输出敏感参数。",
    ]:
        add_bullet(doc, text)

    add_heading(doc, "2.2 客户端必测场景", 2)
    add_table(doc, ["场景", "预期事件", "关键断言"], [
        ("正常胜利", "run_start … wave_end … run_end", "run_end.result=victory，恰好一次"),
        ("普通怪漏怪失败", "enemy_goal → heart_lost → death_wave → run_end", "顺序正确、heart_after=0"),
        ("Boss 到达终点", "boss_goal → death_wave → run_end", "Boss 立即失败原因明确"),
        ("断线/超时", "run_end", "reason=disconnect_timeout 或 reconnect_timeout"),
        ("符文门槛拒绝", "rune_gate_rejection", "account_day 与稳定原因码存在"),
        ("服务端拒绝交易", "ledger_result", "status=rejected，客户端资产不变"),
        ("Firebase 初始化失败", "无业务副作用", "游戏继续，WriteErrorCount 增加"),
        ("多 Adapter 同局", "连续 sequence", "无 OutOfOrder、无重复 event_id"),
    ], [2100, 3250, 4010], font_size=8.4)

    add_page_break(doc)
    add_heading(doc, "3. 服务端具体任务步骤", 1)
    server_steps = [
        ("S1 协议与版本", "实现 analytics_contract_version；所有权威响应返回稳定状态码、operation_id、server_event_id、server_timestamp。", "OpenAPI/协议文档"),
        ("S2 幂等与账本", "体力、金币、道具、符文、广告奖励、商店和结算全部经过幂等账本；重复请求返回同一结果。", "幂等测试和账本表"),
        ("S3 对局会话", "验证 run_id、run_seed、build/config version；记录开始、结算、异常终止和重连状态。", "对局审计日志"),
        ("S4 权威响应", "客户端不得推断资产成功；服务端返回 accepted/duplicate/rejected/timeout/restored 及稳定 reason。", "响应契约测试"),
        ("S5 广告闭环（延期）", "当前只冻结未来规则：广告完成不等于奖励；广告功能进入排期后再实现验证、幂等 ledger_result 和 grant。", "当前无开发；未来伪造/重放测试"),
        ("S6 商店闭环", "服务端控制资格、Offer 池、购买、候选项失效；返回幂等结果。", "并发购买测试"),
        ("S7 排行结算", "服务端生成 rank_snapshot/rank_change/settlement_gold；客户端只显示和记录确认结果。", "对账与回滚测试"),
        ("S8 数据审计", "保留 server_event_id、run_id、operation_id、状态、原因和时间；禁止在 Firebase 参数中发送 Token/原始交易号。", "审计查询"),
        ("S9 补偿策略", "客户端未收到响应时可重查 operation_id；服务端返回已完成结果，不重复发奖。", "断线恢复测试"),
        ("S10 监控告警", "监控成功率、超时率、重复率、账本不一致、结算延迟和各 Lane 流量。", "服务端 Dashboard/告警"),
    ]
    add_table(doc, ["步骤", "具体任务", "验收产物"], server_steps, [1500, 5360, 2500], font_size=8.3)

    add_heading(doc, "3.1 服务端响应最小字段", 2)
    add_table(doc, ["字段", "规则", "用途"], [
        ("operation_id", "客户端生成或服务端下发，单次业务操作唯一", "请求/响应关联"),
        ("server_event_id", "服务端生成，全局唯一", "审计和重复排查"),
        ("ledger_status", "accepted / duplicate / rejected / timeout / restored", "统一账本状态"),
        ("reason", "稳定 snake_case 原因码", "失败归因"),
        ("server_timestamp", "UTC ISO-8601", "延迟与顺序核对"),
        ("transaction_ref_hash", "非敏感哈希引用", "跨系统对账"),
        ("authoritative_balance", "操作后的权威余额；不作为 Firebase 必传字段", "客户端校准"),
        ("contract_version", "例如 analytics-v3", "灰度兼容"),
    ], [1900, 4300, 3160], font_size=8.5)
    add_callout(doc, "安全要求", "Firebase API Secret、服务端密钥、支付凭据和原始交易 ID 只能存在服务端安全环境，不得进入 Unity 包、google-services.json 自定义字段或埋点参数。", fill=RED)

    add_page_break(doc)
    add_heading(doc, "4. Firebase 具体任务步骤", 1)
    firebase_steps = [
        ("F1 项目与 App", "确认 dragonbound-5f868、Android 包名 com.drakeforge.mergedefense；按 dev/qa/staging/prod 决定独立 Firebase 项目或明确 Lane 隔离。", "项目/App 清单"),
        ("F2 Analytics/GA4", "启用 Analytics；确认 GA4 Property、时区、货币、数据流和访问角色。", "GA4 数据流可用"),
        ("F3 DebugView", "配置开发设备调试；验证参数、事件顺序和 Lane。", "联调截图/记录"),
        ("F4 自定义定义", "为 build_lane、execution_context、side、reason、boss_id、hero_id 等当前高价值参数注册维度；广告参数暂不注册。", "维度字典"),
        ("F5 Key Events", "将 run_end、首次 hero_formed、首个 boss_kill 等业务关键事件按产品需求标记；避免把高频事件全部设为关键事件。", "关键事件清单"),
        ("F6 探索报表", "建立对局漏斗、波次留存、Boss 漏斗、英雄形成、失败原因、经济结果和数据质量探索。", "GA4 Explore 模板"),
        ("F7 BigQuery", "建议启用原始事件导出；建立 production 过滤视图、事件展平视图、日增量 KPI 表。", "数据集和视图"),
        ("F8 数据治理", "访问最小权限、数据保留、删除流程、同意状态、PII 检查、开发数据隔离。", "权限与合规记录"),
        ("F9 告警", "对 run_end 缺失、事件量突降、invalid/duplicate、Lane 污染、结算不一致建立告警。", "质量监控"),
        ("F10 发布观察", "灰度期间按 build_version/config_version/Lane 对比；稳定后固化基线。", "7 天观察报告"),
    ]
    add_table(doc, ["步骤", "具体任务", "验收产物"], firebase_steps, [1500, 5360, 2500], font_size=8.3)

    add_heading(doc, "4.1 推荐 Firebase/GA4 报表", 2)
    for text in [
        "Run Funnel：session_start → run_start → wave 3/6/12/16/20 → run_end。",
        "Boss Funnel：boss_spawn → boss_skill → boss_kill 或 boss_goal。",
        "Hero Funnel：recruit_result → hero_formed → hero_level_up。",
        "Failure Explorer：run_end.reason × wave × build_version × config_version。",
        "Economy Quality：ledger_status × operation × reason；与服务端账本抽样对账。",
        "Data Quality：每局起止完整率、波次配对率、重复率、无效率、Lane 污染率。",
    ]:
        add_bullet(doc, text)

    add_page_break(doc)
    add_heading(doc, "5. 三方必须对齐的完整策划", 1)
    add_heading(doc, "5.1 命名与版本决策", 2)
    add_table(doc, ["概念", "V3 统一名称", "当前 V2 名称", "迁移规则"], [
        ("对局开始", "run_start", "run_start", "保持不变"),
        ("对局结束", "run_end", "match_finish", "V2 映射一个版本；禁止双计 KPI"),
        ("波次开始", "wave_start", "wave_start", "保持不变"),
        ("波次结束", "wave_end", "wave_finish", "V2 映射一个版本；禁止双计 KPI"),
        ("英雄形成", "hero_formed", "hero_formed", "保持不变"),
        ("协议版本", "analytics_contract_version", "event_version", "V3 增加显式契约版本"),
        ("发布通道", "build_lane", "无", "加入所有事件公共字段"),
    ], [1800, 1900, 1900, 3760], font_size=8.4)
    add_callout(doc, "推荐决策", "采用 run_end / wave_end 作为产品与数据统一口径。客户端内部可在迁移期把 V2 match_finish / wave_finish 转换为 V3 名称，但 Firebase 生产流只发送一个名称。", fill=GREEN)

    add_heading(doc, "5.2 公共事件字段", 2)
    add_table(doc, ["字段", "必填", "定义/约束"], [
        ("event_name", "是", "注册事件名，snake_case"),
        ("event_version", "是", "事件结构版本，建议 3"),
        ("analytics_contract_version", "是", "三方契约版本，例如 analytics-v3"),
        ("event_id", "是", "单事件唯一，重试保持不变"),
        ("client_timestamp", "客户端事件", "UTC ISO-8601；顺序以 sequence 为准"),
        ("server_timestamp", "权威结果", "UTC ISO-8601"),
        ("run_id", "对局事件", "每局唯一，不包含账号或设备信息"),
        ("run_seed", "对局事件", "实际模拟种子"),
        ("sequence", "是", "同一 run_id 从 1 连续递增"),
        ("build_lane", "是", "dev / qa / staging / production"),
        ("execution_context", "是", "live_player_vs_ai / diagnostic_ai_vs_ai / hero_slice_showcase"),
        ("build_version", "是", "客户端构建版本"),
        ("config_version", "是", "服务器/内容配置快照版本"),
        ("side", "是", "player / ai / system"),
        ("wave", "是", "未开波为 0，否则 1-20"),
        ("operation_id", "按需", "服务请求与结果关联键"),
        ("reason", "失败/终止时", "稳定原因码，禁止本地化文案"),
    ], [2100, 1100, 6160], font_size=8.2)

    add_heading(doc, "5.3 Build Lane 定义", 2)
    add_table(doc, ["Lane", "用途", "数据策略", "允许用户"], [
        ("dev", "本地开发与编辑器验证", "默认排除所有业务 KPI；允许固定种子", "研发"),
        ("qa", "测试包、自动化与回归", "独立筛选；允许测试账号和诊断上下文", "QA/研发"),
        ("staging", "准生产后端与灰度前验收", "结构应与生产一致；禁止污染 production", "内部验收"),
        ("production", "正式发布", "业务 KPI 唯一默认数据源", "真实用户"),
    ], [1300, 2600, 3360, 2100], font_size=8.5)
    add_para(doc, "Build Lane 与 execution_context 是两个独立维度：production Lane 仍可能有 live_player_vs_ai；qa Lane 可以运行 diagnostic_ai_vs_ai。任何同一 run_id 内 Lane 或 Context 改变都视为非法事件。")

    add_heading(doc, "5.4 广告埋点延期与预留策略", 2)
    add_callout(doc, "当前决策", "广告数据目前没有使用需求。现阶段只冻结事件协议、最小字段、稳定广告位 ID 规则和可选观察接口；不接 Firebase、不发送广告事件、不维护每日/每局计数，也不配置广告报表。", fill=GREEN)
    add_table(doc, ["内容", "现在", "未来触发条件", "说明"], [
        ("事件协议", "定义并冻结", "立即", "ad_request / ad_impression / ad_result / ad_reward_result"),
        ("广告位 ID 规则", "定义", "广告位进入产品规划", "使用稳定 snake_case；文案变化不改 ID"),
        ("观察接口", "按业务存在情况", "已有或近期接入 AdService", "IAdAnalyticsObserver + NoOp；业务不得直接调用 Firebase"),
        ("Firebase Adapter", "不做", "开始需要广告数据或进入商业化 QA", "未来实现 FirebaseAdAnalyticsAdapter 并替换 NoOp"),
        ("Firebase 自定义维度", "不做", "开始建立广告报表", "只注册高价值、低基数字段"),
        ("客户端频次计数", "不做", "原则上不需要", "从原始事件按用户/session/run 聚合计算"),
        ("持久化埋点队列", "不为广告单独做", "统一 Analytics 可靠性阶段", "复用统一队列，不创建广告专属队列"),
        ("服务端奖励验证", "延期", "广告奖励功能进入开发", "届时实现验证、幂等、账本与重复领取保护"),
    ], [1800, 1500, 3000, 3060], font_size=7.9)

    add_heading(doc, "5.4.1 广告事件与最小字段", 3)
    add_table(doc, ["事件", "未来触发点", "最小专用字段"], [
        ("ad_request", "客户端向广告 SDK 请求展示", "ad_operation_id, ad_point_id, ad_format"),
        ("ad_impression", "广告 SDK 确认真实展示", "ad_operation_id, ad_point_id, ad_format"),
        ("ad_result", "SDK 返回完成、关闭或失败", "ad_operation_id, result, reason"),
        ("ad_reward_result", "服务端确认奖励发放或拒绝", "ad_operation_id, result, reason, reward_type, reward_amount"),
    ], [1800, 3300, 4260], font_size=8.3)
    add_para(doc, "公共字段继续复用 event_id、event_version、run_id/session_id、sequence、build_lane、build_version、config_version 和时间戳。现在不增加 daily_impression_index、session_impression_index、run_impression_index、daily_ad_count 或 total_ad_count；这些均属于可由 Firebase/BigQuery 后期计算的派生指标。")

    add_heading(doc, "5.4.2 可选观察接口", 3)
    add_para(doc, "如果广告业务已经存在或近期会建立统一 AdService，则同步增加与 Firebase 无关的 IAdAnalyticsObserver，并默认注入 NoOpAdAnalyticsObserver。未来需要数据时新增 FirebaseAdAnalyticsAdapter，无需修改广告业务调用点。如果广告功能尚不存在且排期未定，则暂不创建代码接口，只在本任务书中保留协议。")
    add_para(doc, "建议基础结果码：completed、closed_early、not_available、load_failed、show_failed、timeout、frequency_capped、consent_denied、sdk_error。服务端奖励结果码：accepted、duplicate、rejected、verification_failed、reward_already_granted。已有原因码只允许新增，不得修改历史含义。")
    add_callout(doc, "启动条件", "当广告功能进入当前/下一迭代、需要制定频控、准备 A/B 测试、开始正式 SDK QA 或需要观看/完成/到账数据时，才启动 Firebase Adapter、服务端奖励验证、DebugView 和报表任务。", fill=GOLD)

    add_page_break(doc)
    add_heading(doc, "6. 核心事件定义", 1)
    core_events = [
        ("run_start", "运行上下文与种子被接受后，进入第一波前", "客户端", "run_id, run_seed, build_lane, context", "每局恰好一次且 sequence=1"),
        ("run_end", "胜负、投降、超时或不可恢复终止被提交", "客户端；结算字段来自服务端", "result, reason, elapsed_seconds, final_wave", "每局最多一次；终局事件最后"),
        ("wave_start", "调度器真正开始第 N 波", "客户端", "wave", "同一波最多一次"),
        ("wave_end", "第 N 波生成窗口关闭/规则定义的波次结束", "客户端", "wave, result, residual_enemy_count", "必须与 wave_start 配对"),
        ("hero_formed", "两个组件成功解析并提交为 Hero", "客户端", "hero_id, hero_runtime_id_hash, formation_source", "提交成功后发；拒绝不发"),
        ("recruit_result", "招募批次正式 Commit", "客户端", "recruitment_number, component/basic/forge counts", "总数=5；预览不发"),
        ("formation_snapshot", "批准的关键书签", "客户端", "snapshot_reason 与聚合数量", "禁止逐帧上报"),
        ("heart_lost", "普通敌人漏怪导致心数扣减", "客户端", "count, reason, heart_before/after", "Boss 直接到点单独用 boss_goal"),
        ("death_wave", "一方到达失败条件", "客户端", "wave, reason", "必须先于 defeat run_end"),
        ("boss_spawn", "Boss 实例进入路径", "客户端", "boss_id, speed, max_hp", "每个 Boss 实例一次"),
        ("boss_kill", "Boss 死亡结算完成", "客户端", "boss_id, duration_seconds", "先于 victory run_end"),
        ("boss_goal", "Boss 到达终点并触发立即失败", "客户端", "boss_id, heart_after=0", "先于 death_wave/run_end"),
        ("ledger_result", "服务端账本响应已确认", "客户端记录服务端结果", "operation, status, ref_hash, reason", "服务端权威；幂等"),
        ("settlement_gold", "服务端确认对局金币结算", "客户端记录服务端结果", "gold_amount, reason, operation_id", "不得本地推断"),
    ]
    add_table(doc, ["事件", "精确触发点", "主上报方", "关键字段", "约束"], core_events,
              [1400, 2600, 1500, 2160, 1700], font_size=7.7)

    add_heading(doc, "6.1 run_end 定义", 2)
    add_table(doc, ["字段", "可选值/规则"], [
        ("result", "victory / defeat / abandoned / technical_abort"),
        ("reason", "使用第 7 节稳定原因码"),
        ("final_wave", "最后进入或完成的波次；1-20"),
        ("elapsed_seconds", "从 run_start 到终局提交的单调时长"),
        ("settlement_state", "pending / confirmed / rejected；不等同于胜负"),
        ("dedupe", "event_id = run_id:run_end；每局最多一条"),
    ], [2400, 6960], font_size=8.7)

    add_heading(doc, "6.2 wave_end 定义", 2)
    add_table(doc, ["字段", "可选值/规则"], [
        ("result", "completed / side_defeated / run_terminated / debug_skipped"),
        ("residual_enemy_count", "波次窗口关闭时仍存活、可跨波的敌人数量"),
        ("spawned_count", "本波实际生成数量"),
        ("killed_count / goal_count", "本波归属的聚合结果；不可用时保持缺省"),
        ("duration_seconds", "该波调度窗口时长"),
        ("dedupe", "event_id = run_id:wave:N:end"),
    ], [2400, 6960], font_size=8.7)

    add_heading(doc, "6.3 hero_formed 定义", 2)
    add_table(doc, ["字段", "规则"], [
        ("hero_id", "稳定配置 ID，不使用显示名称"),
        ("hero_runtime_id_hash", "仅用于局内关联；不可反推用户身份"),
        ("recipe_id", "稳定配方 ID"),
        ("formation_source", "post_drop / recruit_auto / debug；生产 KPI 排除 debug"),
        ("component_count", "实际消耗组件数"),
        ("elapsed_seconds", "从 run_start 到形成提交的时长"),
        ("dedupe", "run_id:side:hero_runtime_id_hash:formed"),
    ], [2400, 6960], font_size=8.7)

    add_page_break(doc)
    add_heading(doc, "7. 失败原因与状态码体系", 1)
    add_para(doc, "所有原因码使用小写 snake_case，必须稳定、可枚举、不可包含本地化文本或动态异常消息。展示文案由客户端根据原因码映射。新增原因码必须走契约评审。")
    add_heading(doc, "7.1 run_end 原因码", 2)
    add_table(doc, ["原因码", "适用 result", "定义"], [
        ("victory_all_waves", "victory", "完成全部 20 波并提交胜利"),
        ("defeat_heart_zero_normal", "defeat", "普通敌人漏怪导致心数归零"),
        ("defeat_boss_goal", "defeat", "Boss 到达终点触发立即失败"),
        ("player_surrender", "abandoned", "玩家主动投降"),
        ("afk_timeout", "abandoned", "达到 AFK 超时时间"),
        ("disconnect_timeout", "technical_abort", "重连宽限结束仍未恢复"),
        ("server_run_rejected", "technical_abort", "服务端拒绝创建或恢复对局"),
        ("client_fatal_error", "technical_abort", "客户端不可恢复异常；只发稳定码"),
        ("app_terminated_unconfirmed", "technical_abort", "下次启动补记的未确认终止"),
        ("debug_abort", "technical_abort", "仅 dev/qa；生产禁止"),
    ], [2700, 1900, 4760], font_size=8.4)

    add_heading(doc, "7.2 操作和服务端原因码", 2)
    add_table(doc, ["领域", "建议原因码"], [
        ("通用", "accepted, duplicate, invalid_request, timeout, offline, server_error, version_mismatch"),
        ("体力", "insufficient_energy, spend_rejected, grant_cap_reached"),
        ("招募", "insufficient_currency, destination_full, invalid_batch, commit_conflict"),
        ("道具", "unknown_item, not_owned, not_equipped, cooldown_active, invalid_target"),
        ("符文", "rune_system_locked_until_day3, unknown_rune, insufficient_fragments, insufficient_owned_copies, loadout_locked"),
        ("广告", "not_completed, verification_failed, reward_already_granted, placement_unavailable"),
        ("商店", "merchant_closed, offer_expired, offer_not_eligible, insufficient_currency, purchase_conflict"),
        ("排行/结算", "settlement_pending, settlement_rejected, rank_snapshot_unavailable, season_closed"),
    ], [2100, 7260], font_size=8.5)

    add_heading(doc, "7.3 状态码与原因码关系", 2)
    add_para(doc, "status 表示处理状态，reason 表示原因。accepted 通常 reason 为空；duplicate 表示同一幂等操作已完成；rejected 必须有业务原因；timeout 表示结果未知，客户端不得自行补发资产；restored 表示重查后恢复出既有结果。")

    add_page_break(doc)
    add_heading(doc, "8. 关键 KPI 规划", 1)
    kpis = [
        ("Run 启动数", "COUNT(run_start)", "production + live context", "客户端"),
        ("Run 完成率", "有 run_end 的 run_id / run_start 的 run_id", "排除 debug_abort", "客户端/Firebase"),
        ("胜率", "run_end.result=victory / 有效 run_end", "按版本、配置、难度分层", "客户端"),
        ("技术中断率", "technical_abort / run_start", "按 reason、设备、版本分层", "客户端"),
        ("波次到达率", "到达 wave N 的 run_id / run_start", "N=3/6/12/16/20", "客户端"),
        ("死亡波次分布", "COUNT(death_wave) BY wave/reason", "区分 Boss 与普通漏怪", "客户端"),
        ("Boss 到达率", "boss_spawn / run_start", "按 boss_id", "客户端"),
        ("Boss 击杀率", "boss_kill / boss_spawn", "按 boss_id、难度", "客户端"),
        ("Boss TTK", "boss_kill.duration_seconds 的 P50/P90", "仅 kill 样本", "客户端"),
        ("英雄形成率", "有 hero_formed 的 run_id / run_start", "排除 debug source", "客户端"),
        ("首次英雄形成时长", "MIN(hero_formed.elapsed_seconds) / run", "P50/P90", "客户端"),
        ("招募转化", "hero_formed / recruit_result", "按 recruitment_number", "客户端"),
        ("符文操作成功率", "operation_result=accepted / 符文操作", "按 operation/reason", "客户端"),
        ("账本成功率", "ledger_status=accepted / ledger_result", "按 operation", "服务端确认"),
        ("结算到账率", "settlement_gold confirmed / 有效 run_end", "允许延迟窗口", "服务端确认"),
        ("事件完整率", "具备 run_start+run_end 的 run_id / run_start", "数据质量 KPI", "Firebase/BigQuery"),
        ("波次配对率", "成对 wave_start/wave_end / wave_start", "数据质量 KPI", "BigQuery"),
        ("无效率", "Recorder Invalid / 尝试记录总数", "目标 <0.1%", "客户端诊断"),
        ("重复率", "Duplicate / 尝试记录总数", "目标 <0.5%", "客户端/服务端"),
    ]
    add_table(doc, ["KPI", "公式", "过滤/分层", "权威来源"], kpis, [1850, 3250, 2650, 1610], font_size=7.8)
    add_callout(doc, "KPI 口径", "所有正式 KPI 默认过滤 build_lane=production 且 execution_context=live_player_vs_ai。任何诊断、展示、测试数据不得混入正式玩家漏斗。", fill=GOLD)

    add_heading(doc, "8.1 首批产品看板", 2)
    for text in [
        "核心健康：DAU、run_start、完成率、胜率、技术中断率。",
        "难度曲线：各波到达率、死亡波次、失败原因、Boss 击杀率和 TTK。",
        "构筑成长：招募次数、英雄形成率、首次形成时长、等级提升、符文使用。",
        "经济闭环：体力消耗、结算金币、奖励与账本成功率、重复/拒绝原因。",
        "数据质量：终局缺失、波次不配对、序列断裂、未知原因码、Lane 污染。",
    ]:
        add_bullet(doc, text)

    add_page_break(doc)
    add_heading(doc, "9. 三方协作、RACI 与交付节奏", 1)
    add_table(doc, ["工作项", "产品/数据", "客户端", "服务端", "Firebase/数据平台", "QA"], [
        ("事件字典与 KPI", "A/R", "C", "C", "C", "C"),
        ("客户端 SDK/Sink", "C", "A/R", "I", "C", "C"),
        ("玩法触发点", "C", "A/R", "I", "I", "C"),
        ("账本与权威结果", "C", "C", "A/R", "I", "C"),
        ("Firebase 配置", "C", "C", "I", "A/R", "C"),
        ("报表与 KPI SQL", "A", "C", "C", "R", "C"),
        ("联调与发布验收", "A", "R", "R", "R", "R"),
        ("隐私与权限", "A", "R", "R", "R", "C"),
    ], [2200, 1450, 1450, 1450, 1650, 1160], font_size=8.1)
    add_para(doc, "R=执行负责，A=最终负责，C=参与评审，I=知会。每个事件必须有且只有一个 A 和一个主上报 Owner。")

    add_heading(doc, "9.1 客户端与服务端连接前后实施原则", 2)
    add_callout(doc, "最优顺序", "连接前冻结三方协议并完成客户端埋点基座；连接过程中按业务垂直切片同步接入权威字段和埋点；连接后完成 staging 对账、Firebase 报表和 production 灰度。不得把全部埋点推迟到接口完成之后。", fill=GREEN)
    add_table(doc, ["阶段", "必须完成", "可并行事项", "阶段 Gate"], [
        ("连接前：协议先行", "V3 事件字典、字段、原因码、KPI、Build Lane、主上报归属", "Firebase 项目/Android 配置；服务端幂等与审计设计", "三方协议签字冻结"),
        ("连接前：客户端基座", "Firebase 初始化、AnalyticsRunSession、统一 sequence、Sink、NoOp、类型化观察口", "客户端玩法事件与模拟服务端响应测试", "InMemory/DebugView 链路通过"),
        ("连接中：垂直切片", "每个 API 同步完成请求、服务端业务、权威字段、客户端适配、埋点与测试", "按启动、过程、终局、资产、排行顺序推进", "切片端到端验收后再进入下一项"),
        ("连接后：全链路", "run_id/operation_id 对账、断线恢复、重复请求、账本一致性", "GA4 探索、BigQuery 视图、质量告警", "staging 全链路通过"),
        ("发布：灰度", "production 小流量、质量指标、回滚开关", "关闭 V2 兼容映射", "连续 7 天达到发布阈值"),
    ], [1850, 3400, 2500, 1610], font_size=7.8)

    add_heading(doc, "9.2 垂直切片排期", 2)
    add_table(doc, ["顺序", "切片", "客户端任务", "服务端任务", "核心事件/验收"], [
        ("1", "对局启动", "生成 operation_id；请求扣体力；成功后创建 Run Session", "幂等扣体力；返回状态、原因、server_event_id", "ledger_result → energy_spend → run_start"),
        ("2", "对局过程", "接入 run/wave、敌人、Boss、招募、阵型、英雄", "提供必要配置版本和对局审计字段", "波次配对、失败顺序、核心漏斗完整"),
        ("3", "终局结算", "产生 run_end；请求金币/段位结算；记录确认结果", "幂等结算；返回金币、段位、账本状态", "run_end + settlement_gold + rank_change + ledger_result"),
        ("4", "道具与符文", "接入 grant/equip/use/craft/reward 观察口", "权威资产、重复保护、拒绝原因", "成功/拒绝均可关联并对账"),
        ("5", "商店与排行", "封装 UI 请求与服务端结果适配", "Offer、购买、排行榜与赛季权威结果", "merchant_* / leaderboard_snapshot"),
        ("6", "广告（延期）", "保持协议和可选 NoOp 接口", "暂不实现奖励验证", "商业化进入排期后单独启用"),
    ], [750, 1550, 2600, 2600, 1860], font_size=7.5)

    add_heading(doc, "9.3 六周建议排期", 2)
    for text in [
        "第 1 周（连接前）：冻结 V3 协议、Build Lane、原因码、核心 KPI 和单一上报归属；完成 Firebase 项目与 Android 配置。",
        "第 2 周（连接前/并行）：客户端完成初始化、统一 Session、Firebase Sink、run/wave 生命周期；服务端完成 operation_id、幂等、账本和审计契约。",
        "第 3 周（连接中）：完成切片 1 对局启动和切片 2 对局过程；接入招募、英雄、敌人、Boss，并执行客户端 QA。",
        "第 4 周（连接中）：完成切片 3 终局结算、切片 4 道具/符文和切片 5 商店/排行；广告继续保持协议/NoOp 预留。",
        "第 5 周（连接后）：完成 run_id/operation_id/账本对账、断网与重复请求测试、GA4 探索、BigQuery 视图、质量告警和 staging 验收。",
        "第 6 周（发布）：production 小流量灰度、连续 7 天稳定观察、验证回滚开关，并在指标达标后关闭 V2 兼容映射。",
    ]:
        add_number(doc, text)

    add_heading(doc, "9.4 排期执行约束", 2)
    for text in [
        "任何服务端接口不得在缺少 operation_id、稳定状态码和原因码的情况下进入联调完成状态。",
        "每个垂直切片必须同时交付业务、埋点、自动化测试和 DebugView/审计证据，不能把埋点作为尾项补做。",
        "客户端玩法事件可在服务端连接前完成；资产、结算、交易和排行只能在服务端权威响应接入后验收。",
        "完整 KPI 报表和账本对账属于连接后任务，不应阻塞连接前的客户端基座建设。",
        "广告不进入当前六周关键路径；只有商业化任务明确排期时才替换 NoOp 并追加服务端奖励验证。",
    ]:
        add_bullet(doc, text)

    add_heading(doc, "9.5 变更流程", 2)
    for text in [
        "新增/删除/改名事件必须提升 contract version，并由产品、客户端、服务端、数据共同评审。",
        "字段从可选改为必填属于破坏性变更；必须提供灰度兼容期。",
        "原因码只能新增，不能重用旧含义；废弃码保留历史解释。",
        "Firebase 自定义维度数量有限，应优先注册高价值、低基数参数。",
        "生产 Schema 变更必须先在 qa/staging Lane 验证，再进入 production。",
    ]:
        add_bullet(doc, text)

    add_page_break(doc)
    add_heading(doc, "10. 联调、验收与发布清单", 1)
    add_heading(doc, "10.1 三方联调顺序", 2)
    for text in [
        "客户端使用 dev Lane，在 Firebase DebugView 验证公共字段、事件名和参数类型。",
        "服务端提供 staging 权威响应，验证 operation_id、状态和原因码。",
        "客户端执行完整一局，导出事件链并与服务端审计日志按 run_id/operation_id 对齐。",
        "数据侧在 BigQuery/GA4 复算核心 KPI，与客户端测试期望和服务端账本核对。",
        "QA 执行断网、重连、重复请求、切后台、崩溃、版本不兼容和隐私关闭场景。",
        "production 灰度后检查事件量、终局完整率、Lane 污染和异常原因分布。",
    ]:
        add_number(doc, text)

    add_heading(doc, "10.2 发布 Gate", 2)
    add_table(doc, ["Gate", "通过标准", "证据"], [
        ("Schema Gate", "未知事件/字段/原因码为 0", "契约测试"),
        ("Sequence Gate", "同局连续且无跨 Adapter 冲突", "EditMode + BigQuery 检查"),
        ("Lifecycle Gate", "run 起止完整；wave 起止配对", "全链路测试"),
        ("Authority Gate", "资产变化均有服务端 accepted/duplicate 结果", "账本对账"),
        ("Firebase Gate", "DebugView 与 staging 报表字段正确", "联调记录"),
        ("Privacy Gate", "关闭同意后无 Analytics 采集；无 PII", "隐私审计"),
        ("Reliability Gate", "Firebase 故障不影响玩法；队列有界", "故障注入"),
        ("Build Gate", "Android Release IL2CPP/ARM64 成功", "CI 构建"),
        ("Quality Gate", "终局完整率≥99.5%，无效率<0.1%", "灰度 7 天数据"),
    ], [1900, 4550, 2910], font_size=8.4)

    add_heading(doc, "10.3 回滚方案", 2)
    for text in [
        "客户端提供远程开关，可关闭自定义 Analytics Sink，不影响游戏核心流程。",
        "服务端按 contract_version 兼容上一版本响应，不因 Analytics 字段缺失拒绝业务。",
        "Firebase 报表按 build_version/build_lane 排除异常版本。",
        "事件改名回滚时只恢复映射，不对同一事实双发新旧事件。",
    ]:
        add_bullet(doc, text)

    add_page_break(doc)
    add_heading(doc, "附录 A：完整事件域与实施状态", 1)
    add_table(doc, ["事件域", "事件", "目标状态"], [
        ("Run/Wave", "run_start, run_end, wave_start, wave_end, death_wave, heart_lost", "P1/P2"),
        ("Enemy/Boss", "enemy_spawn, enemy_goal, last_hit, boss_spawn, boss_skill, boss_summon, boss_damage_window, boss_kill, boss_goal", "P2"),
        ("Recruit/Formation", "recruit_result, formation_snapshot, hero_formed, hero_xp, hero_level_up", "P2"),
        ("Items", "item_grant, item_equip, item_use", "P2/P3"),
        ("Runes", "rune_grant, rune_equip, rune_loadout_assign, rune_loadout_unequip, rune_craft, rune_gate_rejection, rune_reward_pending/granted/rejected", "P2/P3"),
        ("Energy", "energy_spend, energy_grant", "P3"),
        ("Ads（暂不采集）", "ad_request, ad_impression, ad_result, ad_reward_result", "仅协议；按商业化排期启用"),
        ("Merchant/Ledger", "merchant_open, merchant_offer, merchant_purchase, ledger_result", "P3"),
        ("Rank/Settlement", "rank_snapshot, rank_change, leaderboard_snapshot, settlement_gold", "P3"),
        ("Rescue", "emergency_save", "功能上线时"),
    ], [1900, 5160, 2300], font_size=8.2)

    add_heading(doc, "附录 B：当前项目差距摘要", 1)
    for text in [
        "V2 已注册 45 个事件，但生产侧仅有内存 Sink，没有 FirebaseAnalyticsSink。",
        "符文业务调用链已支持可选 Adapter 注入，默认未创建生产 Adapter。",
        "DrakeforgeAnalyticsAdapterV1 已具备 W12 Boss、伤害窗口、道具和符文适配方法，但 Bootstrap 尚未创建或 Attach。",
        "run/wave 生命周期没有完整类型化回调；不能通过解析日志替代。",
        "多个 Adapter 独立维护 sequence，正式集成前必须统一到 AnalyticsRunSession。",
        "服务端经济、广告、商店、排行和账本边界尚待实现。",
        "V2 Jsonl 生产 Sink 尚不存在；Firebase 接入后由 Firebase Sink 承担生产发送。",
    ]:
        add_bullet(doc, text)

    add_heading(doc, "附录 C：参考资料", 1)
    refs = [
        "项目事件字典：Docs/DragonBound/Analytics/Drakeforge_Analytics_Event_Dictionary_V2.md",
        "项目接入地图：Docs/DragonBound/Analytics/Drakeforge_Analytics_Integration_Map_V1.md",
        "项目差距审计：Docs/DragonBound/Analytics/Drakeforge_Analytics_Integration_Gaps_V1.md",
        "Firebase Unity 接入：https://firebase.google.com/docs/unity/setup",
        "Firebase Unity 事件：https://firebase.google.com/docs/analytics/unity/events",
        "Firebase Unity 发行说明：https://firebase.google.com/support/release-notes/unity",
    ]
    for ref in refs:
        add_bullet(doc, ref)

    add_callout(doc, "评审决策项", "请三方在实施前确认：① run_end/wave_end 命名迁移；② Build Lane 的 Firebase 项目隔离方式；③服务端权威事件由客户端记录服务端确认，还是未来由服务端 Measurement Protocol 直发；④BigQuery 是否首期启用；⑤用户同意策略与投放地区。", fill=GOLD)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build_doc())
